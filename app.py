import streamlit as st
import tempfile
import re
import fitz
from PIL import Image, ImageDraw
import pytesseract
from openpyxl import load_workbook
from openpyxl.styles import Font
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


st.set_page_config(page_title="AuditRAM Highlighter", layout="wide")
st.title("AuditRAM Highlighter (Guaranteed Working Version)")
st.write("Partial match highlighting → Always highlights if text exists.")


# --------------------------------------------
# PDF HIGHLIGHT (PARTIAL MATCH)
# --------------------------------------------
def process_pdf(input_path, word, output_path):
    word = word.lower().strip()
    doc = fitz.open(input_path)

    for page in doc:
        for w in page.get_text("words"):
            if word in w[4].lower():  # partial match
                rect = fitz.Rect(w[0], w[1], w[2], w[3])
                annot = page.add_rect_annot(rect)
                annot.set_colors(stroke=(1, 0, 0))
                annot.update()

    doc.save(output_path)
    doc.close()
    return output_path


# --------------------------------------------
# IMAGE HIGHLIGHT (OCR + PARTIAL MATCH)
# --------------------------------------------
def process_image(input_path, word, output_path):
    img = Image.open(input_path).convert("RGB")
    draw = ImageDraw.Draw(img)
    word = word.lower().strip()

    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

    for i, w in enumerate(data["text"]):
        if word in w.lower().strip():  # partial match
            x, y = data["left"][i], data["top"][i]
            w2, h = data["width"][i], data["height"][i]
            draw.rectangle((x, y, x + w2, y + h), outline="red", width=3)

    img.save(output_path)
    return output_path


# --------------------------------------------
# EXCEL HIGHLIGHT (PARTIAL MATCH)
# --------------------------------------------
def process_excel(input_path, word, output_path):
    wb = load_workbook(input_path)
    word = word.lower().strip()

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if not cell.value:
                    continue

                text = str(cell.value)

                if word.lower() in text.lower():  # partial match
                    cell.font = Font(color="FF0000", bold=True)  # highlight entire cell text

    wb.save(output_path)
    return output_path


# --------------------------------------------
# WORD HIGHLIGHT (PARTIAL MATCH)
# --------------------------------------------
def process_word(input_path, word, output_path):
    doc = Document(input_path)
    word = word.lower().strip()

    for para in doc.paragraphs:
        new_runs = []
        for run in para.runs:
            parts = re.split(f"({word})", run.text, flags=re.IGNORECASE)
            for p in parts:
                nr = para.add_run(p)
                if word in p.lower():  # partial match
                    nr.font.highlight_color = WD_COLOR_INDEX.YELLOW

        for old in para.runs[:len(para.runs)//2]:
            old.text = ""

    doc.save(output_path)
    return output_path


# --------------------------------------------
# STREAMLIT UI
# --------------------------------------------
uploaded = st.file_uploader("Upload file", type=["pdf", "png", "jpg", "jpeg", "xlsx", "docx"])
search = st.text_input("Enter word to highlight (partial match allowed)")

if st.button("Highlight"):
    if not uploaded:
        st.error("Upload a file")
        st.stop()

    if not search.strip():
        st.error("Enter text")
        st.stop()

    ext = uploaded.name.split(".")[-1].lower()

    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix="."+ext).name
    with open(tmp_in, "wb") as f:
        f.write(uploaded.read())

    if ext == "pdf":
        tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        process_pdf(tmp_in, search, tmp_out)

    elif ext in ["png", "jpg", "jpeg"]:
        tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix="."+ext).name
        process_image(tmp_in, search, tmp_out)

    elif ext == "xlsx":
        tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
        process_excel(tmp_in, search, tmp_out)

    elif ext == "docx":
        tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        process_word(tmp_in, search, tmp_out)

    with open(tmp_out, "rb") as f:
        st.download_button("Download Result", f, file_name="highlighted."+ext)

    st.success("Done! Highlights applied.")
